"""Orchestrates end-to-end document ingestion into all backend stores."""

from __future__ import annotations

import asyncio
import time
from pathlib import Path
from typing import Any, Literal

from .chunker import SmartChunker
from .csv_json_loader import StructuredDataLoader
from .embedder import BaseEmbedder, BGEEmbedder
from .metadata_extractor import MetadataExtractor
from .models import DocumentMetadata, IngestionResult, ProcessedDocument, Section
from .pdf_processor import PDFProcessor
from .web_scraper import WebScraper


class IngestionPipeline:
	"""Main ingestion coordinator with storage fan-out and cache invalidation."""

	def __init__(
		self,
		pdf_processor: PDFProcessor | None = None,
		web_scraper: WebScraper | None = None,
		chunker: SmartChunker | None = None,
		embedder: BaseEmbedder | None = None,
		vector_store: Any | None = None,
		relational_store: Any | None = None,
		bm25_store: Any | None = None,
		graph_store: Any | None = None,
		cache_store: Any | None = None,
	) -> None:
		self.pdf_processor = pdf_processor or PDFProcessor()
		self.web_scraper = web_scraper or WebScraper(self.pdf_processor)
		self.chunker = chunker or SmartChunker()
		self.embedder = embedder or BGEEmbedder()
		self.structured_loader = StructuredDataLoader()
		self.metadata_extractor = MetadataExtractor()

		self.vector_store = vector_store
		self.relational_store = relational_store
		self.bm25_store = bm25_store
		self.graph_store = graph_store
		self.cache_store = cache_store

	async def ingest(
		self,
		source: str | Path | bytes,
		source_type: Literal["pdf", "url", "csv", "json", "docx"],
		metadata_override: dict[str, Any] | None = None,
	) -> IngestionResult:
		"""Ingest one source end-to-end across retrieval backends."""
		t0 = time.perf_counter()
		warnings: list[str] = []

		processed = await self._process_source(source=source, source_type=source_type)
		if metadata_override:
			for key, value in metadata_override.items():
				if hasattr(processed.metadata, key):
					setattr(processed.metadata, key, value)
				else:
					processed.metadata.extra[key] = value

		extracted = self.metadata_extractor.extract_basic(processed.raw_text)
		for k, v in extracted.items():
			if v is not None and getattr(processed.metadata, k, None) in (None, ""):
				setattr(processed.metadata, k, v)

		chunks = self.chunker.chunk_document(processed)
		if chunks:
			vectors = self.embedder.embed_documents([c.text for c in chunks])
			for chunk, emb in zip(chunks, vectors.tolist()):
				chunk.embedding = emb

		entities = self.metadata_extractor.extract_entities(processed.raw_text)
		await self._store_everywhere(processed=processed, chunks=chunks, entities=entities)
		await self._invalidate_cache_for_doc(processed.metadata.doc_id)

		return IngestionResult(
			doc_id=processed.metadata.doc_id,
			source_type=source_type,
			chunks_created=len(chunks),
			entities_extracted=len(entities),
			ingestion_time_ms=(time.perf_counter() - t0) * 1000,
			warnings=warnings,
		)

	async def ingest_batch(self, sources: list[dict[str, Any]]) -> list[IngestionResult]:
		"""Ingest many sources concurrently."""
		tasks = [
			self.ingest(
				source=item["source"],
				source_type=item["source_type"],
				metadata_override=item.get("metadata_override"),
			)
			for item in sources
		]
		return await asyncio.gather(*tasks)

	async def delete_document(self, doc_id: str) -> bool:
		"""Delete a document from all stores and invalidate cache."""
		ok = True
		for store, method in [
			(self.vector_store, "delete"),
			(self.relational_store, "delete_document"),
			(self.bm25_store, "remove_document"),
			(self.graph_store, "delete_document"),
		]:
			if store is None:
				continue
			fn = getattr(store, method, None)
			if fn is None:
				continue
			try:
				result = fn(doc_id)
				if asyncio.iscoroutine(result):
					await result
			except Exception:
				ok = False

		await self._invalidate_cache_for_doc(doc_id)
		return ok

	async def update_document(self, doc_id: str, source: str | Path | bytes) -> IngestionResult:
		"""Replace an existing document with newly ingested content."""
		await self.delete_document(doc_id)
		result = await self.ingest(source=source, source_type=self._infer_source_type(source))
		return result

	async def _process_source(
		self,
		source: str | Path | bytes,
		source_type: Literal["pdf", "url", "csv", "json", "docx"],
	) -> ProcessedDocument:
		if source_type == "pdf":
			return self.pdf_processor.process(source if isinstance(source, (str, Path)) else Path("inline.pdf"))
		if source_type == "url":
			if not isinstance(source, str):
				raise TypeError("URL source must be a string")
			return await self.web_scraper.scrape_url(source)
		if source_type == "csv":
			if not isinstance(source, (str, Path)):
				raise TypeError("CSV source must be a path")
			text = self.structured_loader.load_csv(source)
			return self._from_text(text, source)
		if source_type == "json":
			text = self.structured_loader.load_json(source)
			return self._from_text(text, source)
		if source_type == "docx":
			text = self._load_docx(source)
			return self._from_text(text, source)
		raise ValueError(f"Unsupported source_type: {source_type}")

	def _from_text(self, text: str, source: str | Path | bytes) -> ProcessedDocument:
		src = str(source) if not isinstance(source, bytes) else "inline-bytes"
		metadata = DocumentMetadata(doc_id=f"doc-{abs(hash(src))}", source=src, title=Path(src).stem)
		return ProcessedDocument(
			raw_text=text,
			sections=[Section(name="Document", text=text, page_start=1, page_end=1)],
			metadata=metadata,
		)

	def _load_docx(self, source: str | Path | bytes) -> str:
		try:
			from docx import Document  # type: ignore
		except Exception:
			return ""

		if isinstance(source, bytes):
			tmp = Path("/tmp/ingest_doc.docx")
			tmp.write_bytes(source)
			doc_path = tmp
		else:
			doc_path = Path(source)

		document = Document(str(doc_path))
		return "\n".join(p.text for p in document.paragraphs if p.text.strip())

	async def _store_everywhere(self, processed: ProcessedDocument, chunks: list, entities: list[str]) -> None:
		if self.vector_store and hasattr(self.vector_store, "add"):
			result = self.vector_store.add(chunks)
			if asyncio.iscoroutine(result):
				await result
		if self.relational_store and hasattr(self.relational_store, "upsert_document"):
			result = self.relational_store.upsert_document(processed.metadata, chunks)
			if asyncio.iscoroutine(result):
				await result
		if self.bm25_store and hasattr(self.bm25_store, "add_chunks"):
			result = self.bm25_store.add_chunks(chunks)
			if asyncio.iscoroutine(result):
				await result
		if self.graph_store and hasattr(self.graph_store, "add_document_graph"):
			result = self.graph_store.add_document_graph(processed)
			if asyncio.iscoroutine(result):
				await result
		if self.graph_store and hasattr(self.graph_store, "add_entities"):
			result = self.graph_store.add_entities(processed.metadata.doc_id, entities)
			if asyncio.iscoroutine(result):
				await result

	async def _invalidate_cache_for_doc(self, doc_id: str) -> None:
		if not self.cache_store or not hasattr(self.cache_store, "invalidate_by_doc"):
			return
		result = self.cache_store.invalidate_by_doc(doc_id)
		if asyncio.iscoroutine(result):
			await result

	def _infer_source_type(self, source: str | Path | bytes) -> Literal["pdf", "url", "csv", "json", "docx"]:
		if isinstance(source, bytes):
			return "json"
		s = str(source)
		if s.startswith("http://") or s.startswith("https://"):
			return "url"
		suffix = Path(s).suffix.lower()
		if suffix == ".pdf":
			return "pdf"
		if suffix == ".csv":
			return "csv"
		if suffix == ".docx":
			return "docx"
		return "json"


__all__ = ["IngestionPipeline"]
